#======================================================
# Doc Utility Functions
#======================================================
'''
Version 1.0
Utility functions for writing word docs
'''

# Import modules

import os
import subprocess
import sys

import pandas as pd
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt
from docx.enum.text import WD_LINE_SPACING


#------------------------------
# Utility Functions
#------------------------------
# Create doc table
def make_doc_table(df, save_path, column_names=None, title=None, open=True):
    # Create doc
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri (Body)'
    font.size = Pt(10)

    # Add table title
    p = doc.add_paragraph('', style='No Spacing')
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    sentence = p.add_run(title)
    sentence.font.size = Pt(8.5)
    sentence.font.bold = True
    sentence.font.italic = True

    # Add table
    nrows, ncols = df.shape
    table = doc.add_table(rows=nrows+1, cols=ncols) #Add one more row for the header
    table.style = 'TableGrid'

    # Add header
    hdr_cells = table.rows[0].cells
    if column_names == None:
        column_names = df.columns
    for col_idx, col_name in enumerate(column_names):
        hdr_cells[col_idx].text = col_name
        hdr_cells[col_idx].paragraphs[0].runs[0].font.size = Pt(10)
        #hdr_cells[col_idx].paragraphs[0].runs[0].font.bold = True
        # Black shading
        black_shading = parse_xml(r'<w:shd {} w:fill="000000"/>'.format(nsdecls('w')))
        hdr_cells[col_idx]._tc.get_or_add_tcPr().append(black_shading)

    # Iterate through each row
    for row_idx, df_row in df.iterrows():
        row_idx = row_idx+1
        # Iterate through each col
        for col_idx, col_val in enumerate(df_row):
            table.rows[row_idx].cells[col_idx].text = str(col_val)
            table.rows[row_idx].cells[col_idx].paragraphs[0].runs[0].font.size = Pt(8.5)
            table.rows[row_idx].cells[col_idx].paragraphs[0].runs[0].font.bold = False

    # Save
    doc.save(save_path)

    # Open
    if open:
        if sys.platform.startswith('darwin'):
            subprocess.call(('open', save_path))
        elif os.name == 'nt': # For Windows
            os.startfile(save_path)
        elif os.name == 'posix': # For Linux, Mac, etc.
            subprocess.call(('xdg-open', save_path))
